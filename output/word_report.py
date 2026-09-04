"""Phase 6-8 Stage 2: single-file Word report generation.

Builds one ``.docx`` per pipeline run (design spec section 29: one Word
file per run, not per query -- multiple queries become subsections within
the report's fixed "7. Candidate Ranking" / "8. Candidate Details"
sections instead, see
``claude/phase678_stage2_word_report_investigation.md`` item 1).

This module owns everything ``python-docx``-specific: low-level OOXML
helpers for bookmarks, external hyperlinks, and the Table of Contents
field (``python-docx`` has no high-level API for any of the three), plus
document assembly. It reuses ``output/report_v2.py`` for row
shaping/selection (the same consolidated rows the Excel workbook is built
from -- see ``build_workbook_sheets``) and ``output/word_narrative.py``
for the deterministic "why ranks highly" / "Biological Interpretation"
text. It has no knowledge of Excel/openpyxl -- the reverse direction
(Excel's ``word_report_link`` column) lives in ``output/excel.py`` and
depends on this module's ``bookmark_name``/``select_top_candidates_per_query``
re-exports from ``output/report_v2.py``, not on this module directly, so
the two writers stay independently callable.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

from analysis.scoring_engine_config import ScoringEngineConfig, load_scoring_engine_config
from core.exceptions import WordReportError
from output.report_v2 import (
    TIER_SAFETY_NET,
    bookmark_name,
    build_workbook_sheets,
    select_top_candidates_per_query,
)
from output.word_narrative import CategoryRef, build_biological_interpretation, build_evolutionary_closer, build_why_ranks_highly


# ---------------------------------------------------------------------------
# Low-level OOXML helpers -- python-docx has no high-level API for any of
# these three (bookmarks, external hyperlinks, TOC field), so this is the
# well-documented ~15-line-per-helper escape hatch into raw OOXML that
# claude/phase678_excel_word_redesign_investigation.md (item 3) anticipated.
# All confirmed working (create -> save -> reopen -> structurally verify)
# during this module's own M1 smoke test.
# ---------------------------------------------------------------------------


def _ensure_hyperlink_style(document: Document) -> None:
    """Add a "Hyperlink" character style if the document's template lacks one.

    python-docx's default template does not ship a "Hyperlink" style (only
    Heading 1-9 and a handful of others) -- without this, a hyperlink run
    styled "Hyperlink" would reference a style that does not exist, which
    Word tolerates but renders as plain unstyled text (still clickable,
    just visually indistinguishable from a normal run).
    """
    if "Hyperlink" in document.styles:
        return
    style = document.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
    style.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    style.font.underline = True


def _set_update_fields_on_open(document: Document) -> None:
    """Force Word to recompute fields (the TOC in particular) when the file is opened.

    Without this, python-docx's TOC field shows only its placeholder text
    ("Right-click and choose Update Field...") until a person manually
    updates it once -- the user explicitly asked to avoid requiring that
    manual step. ``<w:updateFields w:val="true"/>`` in settings.xml is the
    standard OOXML mechanism for "recalculate fields on open"; confirmed to
    round-trip correctly (present after save + reopen) during this
    module's M1 smoke test.
    """
    settings = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.insert(0, update_fields)


def _add_bookmark(paragraph: Paragraph, name: str, bookmark_id: int) -> None:
    """Wrap ``paragraph``'s existing content in a named bookmark.

    Word bookmark ids only need to be unique within one document; callers
    pass a running counter. ``name`` must already be Word-legal (see
    ``output/report_v2.py::bookmark_name``) -- not re-validated here.
    """
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_external_hyperlink(paragraph: Paragraph, url: str, text: str) -> None:
    """Append a clickable external hyperlink run to ``paragraph``.

    Not currently called by this module (the Word->Excel reference below
    is deliberately plain text, not a link -- see the module docstring's
    cross-reference to option D in
    claude/phase678_excel_word_redesign_investigation.md item 6), but kept
    here as the validated M1 primitive for a future direction that does
    want a real Word-side external link.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    run_properties.append(style)
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_toc_field(document: Document) -> None:
    """Insert a Word Table of Contents field covering heading levels 1-3.

    Candidate-detail headings are deliberately level 4 (see
    ``_write_candidate_details``) so the TOC lists report sections and
    per-query subsections only, not every individual candidate.
    """
    paragraph = document.add_paragraph()
    begin_run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    begin_run._r.append(fld_begin)
    begin_run._r.append(instr)

    separate_run = paragraph.add_run()
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(fld_separate)

    placeholder_run = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "Right-click and choose “Update Field” to generate the table of contents."
    placeholder_run.append(placeholder_text)
    paragraph._p.append(placeholder_run)

    end_run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)


# ---------------------------------------------------------------------------
# Category references (which row columns feed "why ranks highly" /
# "Biological Interpretation", and what they are called/capped at) --
# resolved once per run from the run's actual config, per scoring_model.
# ---------------------------------------------------------------------------


def category_refs_for_scoring_model(
    scoring_model: str,
    engine_config: ScoringEngineConfig,
    legacy_weights: Any,
) -> tuple[CategoryRef, ...]:
    """Resolve the CategoryRef list output/word_narrative.py should enumerate for one run.

    v2_evidence_based uses the scoring engine's own category_caps (the
    same numbers 05_Sequence_Evidence..10_Negative_Evidence categorize
    Interaction_Evidence_Detail rows by). legacy_additive has no category
    concept at all -- its closest analogues are the fixed
    interaction_scoring.scoring_weights point budget, kept as separate
    line items (e.g. co_occurrence and domain_complementarity are not
    combined the way v2's functional_domain_score already is) since that
    is how legacy_additive actually computes and exposes them.
    """
    if scoring_model == "v2_evidence_based":
        caps = engine_config.category_caps
        interaction_cap = (
            caps.get("external_ppi_evidence", 0.0)
            + caps.get("coexpression_evidence", 0.0)
            + caps.get("pih_direct_interaction", 0.0)
        )
        return (
            CategoryRef("candidate_priority_score", "sequence", "Sequence/Source Classification", caps.get("source_classification", 0.0)),
            CategoryRef("same_gene_neighborhood_score", "genomic_context", "Genomic Context", caps.get("genomic_context", 0.0)),
            CategoryRef("functional_domain_score", "functional_domain", "Functional/Domain", caps.get("functional_annotation", 0.0)),
            CategoryRef("interaction_evidence_score", "interaction", "Interaction", interaction_cap),
            CategoryRef("evolutionary_score", "evolutionary", "Evolutionary", caps.get("pih_evolutionary", 0.0)),
            CategoryRef("cellular_compatibility_score", "cellular_compatibility", "Cellular Compatibility", caps.get("pih_cellular_compatibility", 0.0)),
        )

    weights = legacy_weights
    return (
        CategoryRef("candidate_priority_score", "sequence", "Sequence/Source Classification", getattr(weights, "candidate_priority", 0.0)),
        CategoryRef("same_gene_neighborhood_score", "genomic_context", "Genomic Context", getattr(weights, "gene_neighborhood", 0.0)),
        CategoryRef("co_occurrence_score", "functional_domain", "Co-occurrence", getattr(weights, "co_occurrence", 0.0)),
        CategoryRef("domain_complementarity_score", "functional_domain", "Domain Complementarity", getattr(weights, "domain_complementarity", 0.0)),
        CategoryRef("string_ppi_score", "interaction", "Interaction (STRING PPI)", getattr(weights, "external_ppi", 0.0)),
    )


# ---------------------------------------------------------------------------
# Section 5: Evidence Architecture (5.1-5.7)
# ---------------------------------------------------------------------------


def _write_evidence_architecture(
    document: Document,
    scoring_model: str,
    engine_config: ScoringEngineConfig,
    pih_bundle_configured: bool,
) -> None:
    """Write the fixed-text "5. Evidence Architecture" section (5.1-5.7).

    Text is static except for: the run's scoring_model (stated once, up
    front) and the 5.5/5.6 evolutionary/PIH-bundle branch (reused from
    output/word_narrative.build_evolutionary_closer so the same fact is
    never worded two different ways in one report). See
    claude/phase678_stage2_word_report_investigation.md item 4 for why
    5.5/5.6 ("data not supplied, can become available") and 5.7 ("no
    signal implemented, cannot become available today") are deliberately
    described as two different *kinds* of limitation, not lumped together.
    """
    caps = engine_config.category_caps
    document.add_heading("5. Evidence Architecture", level=1)
    document.add_paragraph(
        f"This run used the {scoring_model!r} scoring model. The seven "
        "evidence categories below are the pipeline's full evidence "
        "vocabulary; which ones actually contributed evidence for any "
        "given candidate depends on what data and configuration were "
        "available for this specific run (see each candidate's own "
        "“why this candidate ranks highly” text in section 8)."
    )

    document.add_heading(f"5.1 Sequence Evidence (cap {caps.get('source_classification', 0.0):.0f})", level=2)
    document.add_paragraph(
        "BLAST-based positive/negative classification and best-hit "
        "identity/coverage/E-value strength. Populated for essentially "
        "every candidate that has any BLAST hit at all -- this is the "
        "pipeline's most consistently available evidence category."
    )

    document.add_heading(f"5.2 Functional/Domain Evidence (cap {caps.get('functional_annotation', 0.0):.0f})", level=2)
    document.add_paragraph(
        "Shared or complementary functional annotation (CDD/Pfam domains, "
        "description terms) between query and candidate. Populated "
        "whenever domain annotation is enabled and available for both "
        "proteins."
    )

    document.add_heading(f"5.3 Genomic Context (cap {caps.get('genomic_context', 0.0):.0f})", level=2)
    document.add_paragraph(
        "Genomic proximity between query and candidate genes, from GFF "
        "coordinates -- used as positive evidence only (a distant "
        "candidate is never penalized for being far away, only not "
        "credited for being close). Populated whenever GFF neighborhood "
        "data is available for both genes."
    )

    interaction_cap = (
        caps.get("external_ppi_evidence", 0.0)
        + caps.get("coexpression_evidence", 0.0)
        + caps.get("pih_direct_interaction", 0.0)
    )
    document.add_heading(f"5.4 Interaction Evidence (cap {interaction_cap:.0f})", level=2)
    document.add_paragraph(
        "External protein-protein interaction evidence from up to three "
        "independent sources: STRING PPI, GEO transcript coexpression, "
        "and the optional ProteinInteractionHunter (PIH) direct-interaction "
        "bridge. Populated whenever the corresponding optional data source "
        "(STRING taxon ID, GEO coexpression, or a PIH evidence bundle) is "
        "configured for the run; each source is independent and any subset "
        "may be available."
    )

    evolutionary_closer = build_evolutionary_closer(scoring_model, pih_bundle_configured)
    document.add_heading(f"5.5 Evolutionary Evidence (cap {caps.get('pih_evolutionary', 0.0):.0f})", level=2)
    document.add_paragraph(
        "Phylogenetic/evolutionary profile consistency between candidate "
        f"and query, sourced entirely from the optional PIH bridge. {evolutionary_closer}"
    )

    document.add_heading(f"5.6 Cellular Compatibility (cap {caps.get('pih_cellular_compatibility', 0.0):.0f})", level=2)
    document.add_paragraph(
        "Subcellular localization / compatibility evidence, also sourced "
        f"from the optional PIH bridge. {evolutionary_closer}"
    )

    document.add_heading("5.7 Negative Evidence (reserved)", level=2)
    document.add_paragraph(
        "Reserved for evidence that directly contradicts a candidate/query "
        "pairing -- e.g. incompatible cellular localization, phylogenetic "
        "inconsistency, or functionally contradictory annotation. No such "
        "signal is implemented in this pipeline version. This category is "
        "reported as not evaluated for every candidate, in every run, with "
        "no exceptions. This is a deliberate design decision, not an "
        "oversight: an earlier implementation attempt used "
        "negative_hit_strength (shown elsewhere in this report, under "
        "candidate_source) as a stand-in for this category and found, on "
        "real-data verification, that it measures a different thing -- how "
        "broadly a candidate protein is conserved across negative "
        "reference genomes -- and using it as a contradiction penalty "
        "incorrectly punished well-conserved true interaction partners. "
        "The two signals are kept visibly separate in this report for "
        "that reason, not merged back together for convenience."
    )


# ---------------------------------------------------------------------------
# Sections 7 / 8: per-query candidate ranking and details
# ---------------------------------------------------------------------------


def _group_by_query(rows: list[dict[str, Any]]) -> list[tuple[str, list[dict[str, Any]]]]:
    """Group already query_id-then-rank-sorted rows into (query_id, rows) pairs.

    Preserves the incoming order (see report_v2.rerank_final_score_rows)
    rather than re-sorting -- query section order is a deliberate,
    reproducible property of the upstream sort, not decided here.
    """
    groups: list[tuple[str, list[dict[str, Any]]]] = []
    current_query: str | None = None
    current_rows: list[dict[str, Any]] = []
    for row in rows:
        query_id = str(row.get("query_id") or "")
        if query_id != current_query:
            if current_query is not None:
                groups.append((current_query, current_rows))
            current_query = query_id
            current_rows = []
        current_rows.append(row)
    if current_query is not None:
        groups.append((current_query, current_rows))
    return groups


def _write_candidate_ranking(document: Document, grouped: list[tuple[str, list[dict[str, Any]]]]) -> None:
    """Write "7. Candidate Ranking": one summary table per query."""
    document.add_heading("7. Candidate Ranking", level=1)
    if not grouped:
        document.add_paragraph("No query-specific candidates were produced by this run.")
        return

    for index, (query_id, rows) in enumerate(grouped, start=1):
        document.add_heading(f"7.{index} Query: {query_id}", level=2)
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        header = table.rows[0].cells
        header[0].text = "Rank"
        header[1].text = "Candidate"
        header[2].text = "Final Score"
        header[3].text = "Tier"
        header[4].text = "Candidate Source"
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = str(row.get("candidate_rank") or "")
            cells[1].text = str(row.get("candidate_protein_id") or "")
            final_score = row.get("final_score")
            cells[2].text = f"{final_score:.1f}" if final_score is not None else "—"
            cells[3].text = str(row.get("final_score_tier") or "—")
            cells[4].text = str(row.get("candidate_source") or "")


def _write_candidate_details(
    document: Document,
    grouped: list[tuple[str, list[dict[str, Any]]]],
    category_refs: tuple[CategoryRef, ...],
    evolutionary_closer: str,
    excel_filename: str,
) -> int:
    """Write "8. Candidate Details": one subsection per candidate, bookmarked for Excel to link to.

    Returns the next unused bookmark id, so callers that add further
    bookmarks afterward do not collide with the ones assigned here.
    """
    document.add_heading("8. Candidate Details", level=1)
    if not grouped:
        document.add_paragraph("No query-specific candidates were produced by this run.")
        return 1

    bookmark_id = 1
    for index, (query_id, rows) in enumerate(grouped, start=1):
        document.add_heading(f"8.{index} Query: {query_id}", level=2)
        n_candidates = len(rows)
        for row in rows:
            candidate_id = str(row.get("candidate_protein_id") or "")
            description = str(row.get("candidate_description") or "").strip()
            title = candidate_id if not description else f"{candidate_id} — {description[:80]}"
            heading = document.add_heading(title, level=4)
            _add_bookmark(heading, bookmark_name(query_id, candidate_id), bookmark_id)
            bookmark_id += 1

            why_paragraph = document.add_paragraph()
            why_paragraph.add_run("Why this candidate ranks highly: ").bold = True
            why_paragraph.add_run(build_why_ranks_highly(row, category_refs))

            interpretation_paragraph = document.add_paragraph()
            interpretation_paragraph.add_run("Biological Interpretation: ").bold = True
            interpretation_paragraph.add_run(
                build_biological_interpretation(
                    row,
                    rank=int(row.get("candidate_rank") or 0),
                    n_candidates=n_candidates,
                    category_refs=category_refs,
                    evolutionary_closer=evolutionary_closer,
                )
            )

            document.add_paragraph(
                f"Full data for this candidate: see {excel_filename}, sheet "
                f"02_Final_Score, query_id={query_id}, "
                f"candidate_protein_id={candidate_id}."
            )
    return bookmark_id


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------


def write_word_report(
    config: Any,
    blast_classification: Any,
    output_path: str | Path,
    interaction_result: Any | None = None,
    excel_filename: str = "",
) -> Path:
    """Write the Phase 6-8 Stage 2 single-file Word report and return its path.

    Mirrors ``output/excel.write_classification_workbook``'s signature
    (same ``config``/``blast_classification``/``interaction_result`` the
    Excel writer already receives at main.py's call site) so both writers
    can be called independently from the same pipeline state.
    ``excel_filename`` is printed as a plain-text cross-reference in each
    candidate's detail section (design spec's Excel cross-link
    requirement, option D -- see
    claude/phase678_excel_word_redesign_investigation.md item 6: reliable
    Word->Excel deep-linking to a specific row is not available, so this
    is a stable filename+id reference rather than a fragile link).
    """
    resolved_output = Path(output_path).expanduser().resolve()
    resolved_output.parent.mkdir(parents=True, exist_ok=True)

    scoring_config = getattr(config, "interaction_scoring", None)
    scoring_model = str(getattr(scoring_config, "scoring_model", "legacy_additive"))
    engine_config = load_scoring_engine_config(getattr(scoring_config, "scoring_engine_config", None))
    legacy_weights = getattr(scoring_config, "scoring_weights", None)
    pih_bundle_configured = bool(getattr(scoring_config, "pih_evidence_bundle", None))
    word_report_config = getattr(scoring_config, "word_report", None)
    max_per_query = int(getattr(word_report_config, "max_candidates_per_query", 15))

    try:
        sheets_data = build_workbook_sheets(config, blast_classification, interaction_result)
        selected_rows = select_top_candidates_per_query(
            sheets_data["final_score_rows"], max_per_query, TIER_SAFETY_NET
        )
        grouped = _group_by_query(selected_rows)
        category_refs = category_refs_for_scoring_model(scoring_model, engine_config, legacy_weights)
        evolutionary_closer = build_evolutionary_closer(scoring_model, pih_bundle_configured)

        document = Document()
        _ensure_hyperlink_style(document)
        _set_update_fields_on_open(document)

        document.add_heading("ProteinHunter Candidate Report", level=0)
        document.add_paragraph(f"Report generated: {datetime.now():%Y-%m-%d %H:%M}")
        document.add_paragraph(f"Scoring model: {scoring_model}")
        document.add_paragraph(
            f"Queries evaluated: {len(grouped)}. Candidates shown per query: up to "
            f"{max_per_query}, plus any additional Tier1_VeryStrong/Tier2_Strong "
            "candidate regardless of rank."
        )
        _add_toc_field(document)

        _write_evidence_architecture(document, scoring_model, engine_config, pih_bundle_configured)
        _write_candidate_ranking(document, grouped)
        _write_candidate_details(document, grouped, category_refs, evolutionary_closer, excel_filename)

        document.save(str(resolved_output))
    except Exception as exc:
        message = (
            f"ProteinHunter could not write the Word report: {resolved_output}. "
            "Please check that the folder is writable and the file is not open."
        )
        raise WordReportError(message) from exc

    return resolved_output


__all__: tuple[str, ...] = (
    "WordReportError",
    "category_refs_for_scoring_model",
    "write_word_report",
)
