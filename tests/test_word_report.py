"""Tests for the Phase 6-8 Stage 2 single-file Word report writer."""

from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from config import (
    INTERACTION_ALPHAFOLD_DEFAULT,
    INTERACTION_EVIDENCE_DETAIL_DEFAULT,
    INTERACTION_NEIGHBORHOOD_DEFAULT,
    INTERACTION_SCORING_WEIGHTS_DEFAULT,
    WORD_REPORT_DEFAULT,
    InteractionScoringConfig,
    WordReportConfig,
)
from output.report_v2 import bookmark_name
from output.word_report import category_refs_for_scoring_model, write_word_report


def blast_classification(**buckets) -> SimpleNamespace:
    defaults = {
        "all_records": {},
        "positive_only_records": {},
        "candidates_relaxed_records": {},
        "positive_all_sources_records": {},
        "negative_unmatched_records": {},
        "no_hit_records": {},
        "negative_hit_records": {},
    }
    defaults.update(buckets)
    return SimpleNamespace(**defaults)


def app_config(
    *,
    scoring_model: str = "v2_evidence_based",
    pih_evidence_bundle: str | None = None,
    word_report: WordReportConfig = WORD_REPORT_DEFAULT,
) -> SimpleNamespace:
    return SimpleNamespace(
        interaction_scoring=InteractionScoringConfig(
            enabled=True,
            query_proteins=(),
            query_fasta=None,
            candidate_sources={"candidates": True},
            max_candidates_per_query=200,
            include_sequences_in_excel=False,
            scoring_weights=INTERACTION_SCORING_WEIGHTS_DEFAULT,
            alphafold=INTERACTION_ALPHAFOLD_DEFAULT,
            neighborhood=INTERACTION_NEIGHBORHOOD_DEFAULT,
            scoring_model=scoring_model,
            scoring_engine_config=None,
            functional_complementarity_ruleset=None,
            pih_evidence_bundle=pih_evidence_bundle,
            evidence_detail_sheet=INTERACTION_EVIDENCE_DETAIL_DEFAULT,
            word_report=word_report,
        )
    )


def _pair_row(query_id: str, candidate_id: str, candidate_source: str = "Candidates", **extra) -> dict:
    base = {
        "query_id": query_id,
        "candidate_protein_id": candidate_id,
        "candidate_source": candidate_source,
        "candidate_description": "a candidate protein",
        "negative_hit_strength": "none",
        "final_score": 10.0,
        "final_score_tier": "Tier4_Weak",
        "evidence_category_count": 1,
        "alphafold_recommended": False,
    }
    base.update(extra)
    return base


def _interaction_result(source_rows: dict[str, list[dict]], scoring_model: str = "v2_evidence_based") -> SimpleNamespace:
    return SimpleNamespace(
        source_rows=source_rows,
        evidence_detail_rows=[],
        evidence_detail_scoring_model=scoring_model,
        query_rows=[{"query_id": qid} for qid in {row["query_id"] for rows in source_rows.values() for row in rows}],
        neighborhood_rows=[],
    )


def test_write_word_report_creates_file(tmp_path: Path) -> None:
    output_path = tmp_path / "reports" / "report.docx"
    source_rows = {"Interaction_Candidates": [_pair_row("q1", "c1", final_score=80.0, final_score_tier="Tier1_VeryStrong")]}

    result = write_word_report(
        config=app_config(),
        blast_classification=blast_classification(),
        output_path=output_path,
        interaction_result=_interaction_result(source_rows),
        excel_filename="results.xlsx",
    )

    assert result == output_path.resolve()
    assert result.exists()
    Document(str(result))  # must open without raising


def test_word_report_has_expected_top_level_sections(tmp_path: Path) -> None:
    output_path = tmp_path / "report.docx"
    source_rows = {"Interaction_Candidates": [_pair_row("q1", "c1")]}

    write_word_report(
        config=app_config(),
        blast_classification=blast_classification(),
        output_path=output_path,
        interaction_result=_interaction_result(source_rows),
    )

    document = Document(str(output_path))
    headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
    assert "5. Evidence Architecture" in headings
    assert "7. Candidate Ranking" in headings
    assert "8. Candidate Details" in headings
    assert "7.1 Query: q1" in headings
    assert "8.1 Query: q1" in headings
    assert any(h.startswith("5.7 Negative Evidence") for h in headings)


def test_word_report_multiple_queries_get_separate_subsections(tmp_path: Path) -> None:
    output_path = tmp_path / "report.docx"
    source_rows = {
        "Interaction_Candidates": [
            _pair_row("q1", "c1"),
            _pair_row("q2", "c1"),
        ]
    }

    write_word_report(
        config=app_config(),
        blast_classification=blast_classification(),
        output_path=output_path,
        interaction_result=_interaction_result(source_rows),
    )

    document = Document(str(output_path))
    headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
    assert "7.1 Query: q1" in headings
    assert "7.2 Query: q2" in headings
    assert "8.1 Query: q1" in headings
    assert "8.2 Query: q2" in headings


def test_word_report_settings_force_field_update_on_open(tmp_path: Path) -> None:
    output_path = tmp_path / "report.docx"
    write_word_report(
        config=app_config(),
        blast_classification=blast_classification(),
        output_path=output_path,
        interaction_result=_interaction_result({"Interaction_Candidates": [_pair_row("q1", "c1")]}),
    )

    document = Document(str(output_path))
    assert "updateFields" in document.settings.element.xml


def test_word_report_toc_field_present(tmp_path: Path) -> None:
    output_path = tmp_path / "report.docx"
    write_word_report(
        config=app_config(),
        blast_classification=blast_classification(),
        output_path=output_path,
        interaction_result=_interaction_result({"Interaction_Candidates": [_pair_row("q1", "c1")]}),
    )

    document = Document(str(output_path))
    assert "TOC" in document.element.xml


def test_word_report_candidate_bookmark_matches_report_v2_bookmark_name(tmp_path: Path) -> None:
    output_path = tmp_path / "report.docx"
    write_word_report(
        config=app_config(),
        blast_classification=blast_classification(),
        output_path=output_path,
        interaction_result=_interaction_result({"Interaction_Candidates": [_pair_row("q1", "cand_1")]}),
    )

    document = Document(str(output_path))
    expected_name = bookmark_name("q1", "cand_1")
    assert f'w:name="{expected_name}"' in document.element.xml


def test_word_report_respects_top_n_and_tier_safety_net(tmp_path: Path) -> None:
    output_path = tmp_path / "report.docx"
    # build_workbook_sheets always reranks by final_score (rerank_final_score_rows),
    # so candidate_rank here is assigned from final_score descending: c1=1 .. c5=5.
    rows = [_pair_row("q1", f"c{i}", final_score=100.0 - i) for i in range(1, 6)]
    # rank 5 (c5) is beyond max_per_query=3, but Tier1 -> must still appear.
    rows[4]["final_score_tier"] = "Tier1_VeryStrong"

    word_report = WordReportConfig(enabled=True, max_candidates_per_query=3)
    write_word_report(
        config=app_config(word_report=word_report),
        blast_classification=blast_classification(),
        output_path=output_path,
        interaction_result=_interaction_result({"Interaction_Candidates": rows}),
    )

    document = Document(str(output_path))
    headings = [p.text for p in document.paragraphs if p.style.name == "Heading 4"]
    # c1..c3 (top 3 by final_score) plus c5 (Tier1 safety net) = 4 candidates shown, c4 excluded.
    assert len(headings) == 4
    assert not any(h.startswith("c4") for h in headings)
    assert any(h.startswith("c5") for h in headings)


def test_word_report_legacy_additive_model_does_not_crash(tmp_path: Path) -> None:
    output_path = tmp_path / "report.docx"
    row = _pair_row(
        "q1",
        "c1",
        candidate_priority_score=20.0,
        same_gene_neighborhood_score=5.0,
        co_occurrence_score=3.0,
        domain_complementarity_score=2.0,
        string_ppi_score=0.0,
    )

    write_word_report(
        config=app_config(scoring_model="legacy_additive"),
        blast_classification=blast_classification(),
        output_path=output_path,
        interaction_result=_interaction_result({"Interaction_Candidates": [row]}, scoring_model="legacy_additive"),
    )

    document = Document(str(output_path))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "legacy_additive scoring model" in text


def test_word_report_evidence_architecture_states_pih_bundle_absence(tmp_path: Path) -> None:
    output_path = tmp_path / "report.docx"
    write_word_report(
        config=app_config(pih_evidence_bundle=None),
        blast_classification=blast_classification(),
        output_path=output_path,
        interaction_result=_interaction_result({"Interaction_Candidates": [_pair_row("q1", "c1")]}),
    )

    document = Document(str(output_path))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "was not supplied for this run" in text


def test_word_report_prints_excel_cross_reference_as_plain_text_not_link(tmp_path: Path) -> None:
    output_path = tmp_path / "report.docx"
    write_word_report(
        config=app_config(),
        blast_classification=blast_classification(),
        output_path=output_path,
        interaction_result=_interaction_result({"Interaction_Candidates": [_pair_row("q1", "cand_1")]}),
        excel_filename="ProteinHunter_results.xlsx",
    )

    document = Document(str(output_path))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "ProteinHunter_results.xlsx" in text
    assert "query_id=q1" in text
    assert "candidate_protein_id=cand_1" in text
    hyperlink_targets = [r.target_ref for r in document.part.rels.values() if r.reltype == RT.HYPERLINK]
    assert "ProteinHunter_results.xlsx" not in hyperlink_targets


def test_category_refs_for_v2_include_all_six_categories() -> None:
    from analysis.scoring_engine_config import load_scoring_engine_config

    engine_config = load_scoring_engine_config(None)
    refs = category_refs_for_scoring_model("v2_evidence_based", engine_config, None)
    kinds = {ref.kind for ref in refs}
    assert kinds == {
        "sequence",
        "genomic_context",
        "functional_domain",
        "interaction",
        "evolutionary",
        "cellular_compatibility",
    }


def test_category_refs_for_legacy_use_scoring_weights() -> None:
    refs = category_refs_for_scoring_model("legacy_additive", None, INTERACTION_SCORING_WEIGHTS_DEFAULT)
    caps = {ref.row_key: ref.cap for ref in refs}
    assert caps["candidate_priority_score"] == INTERACTION_SCORING_WEIGHTS_DEFAULT.candidate_priority
    assert caps["co_occurrence_score"] == INTERACTION_SCORING_WEIGHTS_DEFAULT.co_occurrence
